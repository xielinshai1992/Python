# -*- coding: utf-8 -*-
# author:xls
"""
    该模块用于自动生成word报告,配置数据存放在excel中
"""
import time
import docx
from office_operation import *

if __name__== '__main__':

    word_file = r'报告模板.docx'
    excel_file = r'配置文件.xlsx'
    timestamp1 = time.time()

    #TODO:获取excel配置文件信息
    Texts_Info = read_data_from_excel(excel_file,'文本')
    Images_Info = read_data_from_excel(excel_file, '图片')
    Tables_Info = read_data_from_excel(excel_file, '表格')
    doc = getdocument(word_file)
    Texts_Dict = {}
    Images_Dict = {}
    Tables_Dict = {}
    for i in range(1, len(Texts_Info)):
        Texts_Dict[Texts_Info[i][0]] = Texts_Info[i][1]
    for i in range(1, len(Images_Info)):
        Images_Dict[Images_Info[i][0]] = [Images_Info[i][1], Images_Info[i][2], Images_Info[i][3], Images_Info[i][4]]
    for i in range(1, len(Tables_Info)):
        Tables_Dict[Tables_Info[i][0]] = [Tables_Info[i][1], Tables_Info[i][2]]
    print(Texts_Dict)
    # TODO: 插入文本
    for id_text, text_value in Texts_Dict.items():
        matching_cycle(doc, id_text, text_value)
    # TODO: 插入图片
    ImagesId_Dict = { v[1]:k for k,v in Images_Dict.items()}
    # 匹配图片标识信息在word中的位置
    for image_id in ImagesId_Dict.keys():
        for p_index in range(len(doc.paragraphs)):
            if image_id in doc.paragraphs[p_index].text:
                image_name = ImagesId_Dict[image_id]
                doc.paragraphs[p_index].text = ""
                doc.paragraphs[p_index].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                lastRun_index = len(doc.paragraphs[p_index].runs) - 1
                doc.paragraphs[p_index].runs[lastRun_index].add_picture('png\\'+ image_name,
                                                                        width= docx.shared.Cm(Images_Dict[image_name][2]),
                                                                        height= docx.shared.Cm(Images_Dict[image_name][3]))
                doc.paragraphs[p_index + 1].insert_paragraph_before(text=Images_Dict[image_name][0])#插入图片标题
    # TODO: 插入表格
    TablesId_Dict = {v[1]: k for k, v in Tables_Dict.items()}
    # 匹配表格标识信息在word中的位置
    for table_id in TablesId_Dict.keys():
        for p_index in range(len(doc.paragraphs)):
            if table_id in doc.paragraphs[p_index].text:
                table_name = TablesId_Dict[table_id]
                #标识符文本后插入表格
                insert_table_after_text(doc, 'csv\\'+ table_name, table_id)
                #替换当前paragraph文本和样式
                doc.paragraphs[p_index].text = Tables_Dict[table_name][0]
                #doc.paragraphs[p_index].style = '表题'
    doc.save('报告结果.docx')
    timestamp2 = time.time()
    print("耗时：%f" %(timestamp2-timestamp1))