# -*- coding: utf-8 -*-
# author:xls
"""
    操作office对象类，主要用于读写execl、word，支撑报告自动生成功能
"""
import openpyxl
import csv
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def getdocument(filename):
    '''
    获取word文件的document对象
    :param filename: word文件路径
    :return: document对象
    '''
    return Document(filename)

def get_csv_data(csv_filename):
    '''
    从csv文件中获取list[list[]]对象
    :param csv_filename: csv文件路径
    :return:list[list[]]
    '''
    csv_file = open(csv_filename, encoding='UTF-8')
    reader = csv.reader(csv_file)
    data = list(reader)
    return data

def insert_table_after_para(table, paragraph):
    '''
    向指定的paragraph后插入table数据
    :param table:
    :param paragraph:
    :return:
    '''
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

def insert_table_after_text(document, excel_filename, id_text):
    '''
    向word文件的标识符文本后插入表格
    :param document: word文件关联的document对象
    :param excel_filename: excel文件路径（表格数据）
    :param id_text: 标识符文本
    :return:
    '''
    target_p = None
    # 定位标识符文本所在的段落
    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        if paragraph_text.endswith(id_text):
            target_p = paragraph
            break
    if target_p is not None:
        datas = get_csv_data(excel_filename)
        # print(datas)
        # 获得excel文件的行和列,初始化一个空的table
        col = len(datas[0])
        row = len(datas)
        table = document.add_table(rows=row, cols=col, style=None)
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for row_index in range(row):
            for col_index in range(col):
                table.cell(row_index, col_index).text = str(datas[row_index][col_index])
                table.cell(row_index, col_index).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #单元格内容居中
        table.autofit = True
        insert_table_after_para(table, target_p)

def matching_cycle(doc, id_text, replace_text):
    '''
    循环匹配word标记文本,匹配成功后替换
    :param doc: document对象
    :param id_text: 标记文本
    :param replace_text: 替换文本
    :return:
    '''
    for p_index in range(len(doc.paragraphs)):
        if id_text in doc.paragraphs[p_index].text:
            for run_index in range(len(doc.paragraphs[p_index].runs)):
                if doc.paragraphs[p_index].runs[run_index].underline and doc.paragraphs[p_index].runs[run_index].text == id_text:
                    doc.paragraphs[p_index].runs[run_index].text = str(replace_text)
                    doc.paragraphs[p_index].runs[run_index].underline = False

def matching_onlyone(doc, id_text, replace_text):
    '''
    仅匹配一次word标记文本,匹配成功后替换
    :param doc: document对象
    :param id_text: 标记文本
    :param replace_text: 替换文本
    :return:
    '''
    for p_index in range(len(doc.paragraphs)):
        if id_text in doc.paragraphs[p_index].text:
            for run_index in range(len(doc.paragraphs[p_index].runs)):
                if doc.paragraphs[p_index].runs[run_index].underline and doc.paragraphs[p_index].runs[run_index].text == id.text:
                    doc.paragraphs[p_index].runs[run_index].text = str(replace_text)
                    doc.paragraphs[p_index].runs[run_index].underline = False
                    break #标记文本替换完成后跳出循环 无须遍历后续run对象
            break #当前段落匹配到标记文本后跳出循环 无须遍历后续paragraph对象

def read_data_from_excel(excel_filename, sheet_name):
    '''
    获取excel文本内容,返回一个list[list[]]对象
    :param excel_filename: excel文件路径
    :param sheet_name: sheet名称
    :return:
    '''
    result = []
    wb = openpyxl.load_workbook(excel_filename, data_only=True)
    if sheet_name == '':
        sheet = wb.get_active_sheet()
    else:
        sheet = wb.get_sheet_by_name(sheet_name)
    row_num = sheet.max_row
    col_num = sheet.max_column
    for row_index in range(row_num):
        row_list = []
        for col_index in range(col_num):
            row_list.append(sheet.cell(row=row_index+1, column=col_index+1).value)
        result.append(row_list)
    return result

def write_data_to_excel(excel_filename, sheet_name, location, data):
    '''
    向指定的excel文件写数据
    :param excel_filename: excel文件名
    :param sheet_name: sheet名称
    :param location: 表格位置 example 'A1' 'C3'
    :param data: 数据值
    :return:
    '''
    wb = openpyxl.load_workbook(excel_filename)
    sheet = wb.get_sheet_by_name(sheet_name)
    sheet[location] = data
    wb.save(excel_filename)

